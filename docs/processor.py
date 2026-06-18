import csv
import json
import re
import shutil
import unicodedata
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

ROOT = Path("/workspace")
INPUT = ROOT / "input"
OUTPUT = ROOT / "output"
TEMPLATE = ROOT / "templates" / "blank.docx"


def safe_name(value, fallback="未命名"):
    value = unicodedata.normalize("NFKC", value or "").strip()
    value = re.sub(r'[\\/:*?"<>|\x00-\x1f]', "_", value)
    value = re.sub(r"\s+", " ", value).strip(" .")
    return value[:150] or fallback


def unique_path(path):
    path = Path(path)
    if not path.exists():
        return path
    index = 2
    while True:
        candidate = path.with_name(f"{path.stem}_{index}{path.suffix}")
        if not candidate.exists():
            return candidate
        index += 1


def set_run_font(run, blue=False, size=18):
    run.font.name = "BiauKai"
    run.font.size = Pt(size)
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:eastAsia"), "標楷體")
    rfonts.set(qn("w:ascii"), "BiauKai")
    rfonts.set(qn("w:hAnsi"), "BiauKai")
    if blue:
        run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)


def set_cell_text(cell, text, blue=False):
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(str(text))
    set_run_font(run, blue=blue)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_image_to_cell(cell, image_path):
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    with Image.open(image_path) as image:
        width, height = image.size
    target_width = 9.0 * width / height
    if target_width > 16.2:
        run.add_picture(str(image_path), width=Cm(16.2))
    else:
        run.add_picture(str(image_path), height=Cm(9))


def mark_row_keep_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def remove_table_rows_from(table, start):
    for row in list(table.rows)[start:]:
        table._tbl.remove(row._tr)


def rebuild_template_pages(doc, page_count):
    body = doc.element.body
    children = list(body)
    paragraphs = [child for child in children if child.tag == qn("w:p") and "鑑定報告書現況照片" in "".join(child.itertext())]
    tables = [child for child in children if child.tag == qn("w:tbl")]
    if not paragraphs or not tables:
        raise RuntimeError("Word 範本結構不符")
    first_title = deepcopy(paragraphs[0])
    later_title = deepcopy(paragraphs[1] if len(paragraphs) > 1 else paragraphs[0])
    base_table = deepcopy(tables[0])
    section = body.sectPr
    for child in list(body):
        if child is not section:
            body.remove(child)
    for page in range(page_count):
        body.insert(len(body) - 1, deepcopy(first_title if page == 0 else later_title))
        body.insert(len(body) - 1, deepcopy(base_table))


def build_word(images, output_path):
    doc = Document(TEMPLATE)
    rebuild_template_pages(doc, max(1, (len(images) + 1) // 2))
    for paragraph in doc.paragraphs:
        if "鑑定報告書現況照片" in paragraph.text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run)
    current = 0
    for table in doc.tables:
        for row in table.rows:
            mark_row_keep_together(row)
        for header_row, number_row, number_row_2, image_row in [(0, 1, 2, 3), (5, 6, 7, 8)]:
            if current >= len(images):
                remove_table_rows_from(table, 4 if header_row == 5 else 0)
                break
            image_path = images[current]
            number = current + 1
            set_cell_text(table.rows[header_row].cells[2], "A", blue=True)
            set_cell_text(table.rows[number_row].cells[0], number, blue=True)
            set_cell_text(table.rows[number_row_2].cells[0], number, blue=True)
            add_image_to_cell(table.rows[image_row].cells[0], image_path)
            current += 1
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def generate_results(records_json):
    records = json.loads(records_json)
    if OUTPUT.exists():
        shutil.rmtree(OUTPUT)
    OUTPUT.mkdir(parents=True)
    groups = {}
    pending = 0
    labels = {"unrecognized": "地址無法辨識", "incomplete": "地址不完整", "compound": "複合地址待確認"}
    for record in records:
        source = INPUT / safe_name(record.get("filename", ""))
        if not source.exists():
            continue
        status = record.get("status", "incomplete")
        first = safe_name(record.get("first_level", ""), "地址不完整")
        second = safe_name(record.get("second_level", ""), "地址不完整")
        if status == "success" and record.get("first_level") and record.get("second_level"):
            groups.setdefault((first, second), []).append((record, source))
        else:
            pending += 1
            destination = OUTPUT / "待確認" / labels.get(status, "地址不完整")
            destination.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source, unique_path(destination / source.name))
    report = []
    for (first, second), items in sorted(groups.items()):
        items.sort(key=lambda item: (int(item[0].get("photo_number") or 999999), item[0].get("filename", "")))
        destination = OUTPUT / first / second
        destination.mkdir(parents=True, exist_ok=True)
        word_images = []
        for index, (_, source) in enumerate(items, 1):
            target = unique_path(destination / f"照片{index:02d}_{safe_name(source.name)}")
            shutil.copy2(source, target)
            word_images.append(target)
        build_word(word_images, destination / f"{safe_name(first + second)}_貼照片.docx")
        report.append((first, second, len(items)))
    total = sum(row[2] for row in report) + pending
    with (OUTPUT / "分類報告.csv").open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.writer(stream)
        writer.writerow(["總照片數", total])
        writer.writerow(["成功分類數", total - pending])
        writer.writerow(["待確認數", pending])
        writer.writerow([])
        writer.writerow(["第一層地址", "第二層地址", "照片數量"])
        writer.writerows(report)
    archive_path = ROOT / "results.zip"
    if archive_path.exists():
        archive_path.unlink()
    with zipfile.ZipFile(archive_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in OUTPUT.rglob("*"):
            if path.is_file():
                archive.write(path, path.relative_to(OUTPUT))
    return str(archive_path)
