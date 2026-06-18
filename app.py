#!/usr/bin/env python3
from __future__ import annotations

import csv
import hashlib
import io
import json
import mimetypes
import os
import re
import shutil
import subprocess
import sys
import threading
import time
import unicodedata
import uuid
import zipfile
from copy import deepcopy
from datetime import datetime
from email.parser import BytesParser
from email.policy import default
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import quote, unquote, urlparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image
import pytesseract

try:
    from pillow_heif import register_heif_opener
    register_heif_opener()
except ImportError:
    pass


APP_DIR = Path(__file__).resolve().parent
STATIC_DIR = APP_DIR / "static"
TEMPLATE_DIR = APP_DIR / "templates"
WORK_DIR = Path(os.environ.get("DATA_DIR", "/tmp/address-photo-tool"))
SESSION_DIR = WORK_DIR / "sessions"
BLANK_TEMPLATE = TEMPLATE_DIR / "貼照片_空檔案.docx"
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".heic", ".heif", ".tif", ".tiff", ".bmp", ".webp"}
MAX_UPLOAD_BYTES = int(os.environ.get("MAX_UPLOAD_BYTES", str(500 * 1024 * 1024)))
SESSION_TTL_SECONDS = int(os.environ.get("SESSION_TTL_SECONDS", str(24 * 60 * 60)))

TAIWAN_CITIES = (
    "臺北市", "新北市", "桃園市", "臺中市", "臺南市", "高雄市",
    "基隆市", "新竹市", "嘉義市", "新竹縣", "苗栗縣", "彰化縣",
    "南投縣", "雲林縣", "嘉義縣", "屏東縣", "宜蘭縣", "花蓮縣",
    "臺東縣", "澎湖縣", "金門縣", "連江縣",
)

STATUS_LABELS = {
    "success": "分類完成",
    "unrecognized": "地址無法辨識",
    "incomplete": "地址不完整",
    "compound": "複合地址待確認",
}


def ensure_dirs() -> None:
    for path in (WORK_DIR, SESSION_DIR):
        path.mkdir(parents=True, exist_ok=True)


def cleanup_expired_sessions() -> None:
    if not SESSION_DIR.exists():
        return
    cutoff = time.time() - SESSION_TTL_SECONDS
    for path in SESSION_DIR.iterdir():
        try:
            if path.is_dir() and path.stat().st_mtime < cutoff:
                shutil.rmtree(path)
        except OSError:
            continue


def safe_name(value: str, fallback: str = "未命名") -> str:
    value = unicodedata.normalize("NFKC", value or "").strip()
    value = re.sub(r"[\\/:*?\"<>|\x00-\x1f]", "_", value)
    value = re.sub(r"\s+", " ", value).strip(" .")
    return value[:150] or fallback


def unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem, suffix = path.stem, path.suffix
    index = 2
    while True:
        candidate = path.with_name(f"{stem}_{index}{suffix}")
        if not candidate.exists():
            return candidate
        index += 1


def normalize_address_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text or "")
    text = text.replace("台", "臺")
    text = text.replace("－", "-").replace("–", "-").replace("—", "-")
    text = re.sub(r"[\s_]+", "", text)
    return text


def photo_number(filename: str, text: str = "") -> int | None:
    source = unicodedata.normalize("NFKC", f"{filename} {text}")
    patterns = [r"照片\s*0*(\d+)", r"(?:photo|img)[-_ ]*0*(\d+)"]
    for pattern in patterns:
        match = re.search(pattern, source, re.IGNORECASE)
        if match:
            return int(match.group(1))
    return None


def parse_address(text: str) -> dict:
    normalized = normalize_address_text(text)
    city_pattern = "|".join(map(re.escape, TAIWAN_CITIES))
    city_match = re.search(city_pattern, normalized)
    city = city_match.group(0) if city_match else ""
    after_city = normalized[city_match.end():] if city_match else normalized

    district_match = re.match(r"(.{1,6}?(?:區|鄉|鎮|市))", after_city)
    district = district_match.group(1) if district_match else ""
    after_district = after_city[district_match.end():] if district_match else after_city

    road_match = re.match(r"(.{1,16}?(?:路|街)(?:[一二三四五六七八九十百0-9]+段)?)", after_district)
    road = road_match.group(1) if road_match else ""
    remainder = after_district[road_match.end():] if road_match else after_district
    remainder = re.sub(r"^(?:地址|門牌)[:：]?", "", remainder)

    detail_match = re.match(
        r"((?:\d+巷)?(?:\d+弄)?(?:\d+(?:之\d+)?(?:、\d+(?:之\d+)?)*)號?"
        r"(?:\d+樓(?:之\d+)?)?(?:\d+室)?)",
        remainder,
    )
    detail = detail_match.group(1) if detail_match else ""
    detail = detail.removesuffix("號")

    road_count = len(re.findall(r"(?:路|街)(?:[一二三四五六七八九十百0-9]+段)?", normalized))
    compound_markers = len(re.findall(r"\d+(?:之\d+)?號", remainder)) > 1
    explicit_compound = "、" in detail or compound_markers
    uncertain_compound = road_count > 1 or (explicit_compound and not detail)

    first_level = f"{city}{district}{road}" if city and district and road else ""
    if uncertain_compound:
        status = "compound"
    elif not city and not district and not road and not detail:
        status = "unrecognized"
    elif not (city and district and road and detail):
        status = "incomplete"
    else:
        status = "success"

    confidence = 0
    confidence += 28 if city else 0
    confidence += 24 if district else 0
    confidence += 28 if road else 0
    confidence += 20 if detail else 0
    if status == "compound":
        confidence = min(confidence, 55)

    return {
        "city": city,
        "district": district,
        "road": road,
        "detail": detail,
        "first_level": first_level,
        "second_level": detail,
        "status": status,
        "confidence": confidence,
    }


def run_ocr(path: Path) -> str:
    try:
        with Image.open(path) as image:
            image.load()
            if image.mode not in ("RGB", "L"):
                image = image.convert("RGB")
            return pytesseract.image_to_string(
                image,
                lang=os.environ.get("OCR_LANG", "chi_tra+eng"),
                config="--oem 1 --psm 11",
                timeout=60,
            ).strip()
    except Exception:
        return ""


def parse_multipart(headers, body: bytes) -> list[tuple[str, str, bytes]]:
    content_type = headers.get("Content-Type", "")
    if "multipart/form-data" not in content_type:
        raise ValueError("請使用 multipart/form-data 上傳")
    message = BytesParser(policy=default).parsebytes(
        f"Content-Type: {content_type}\r\nMIME-Version: 1.0\r\n\r\n".encode() + body
    )
    files = []
    for part in message.iter_parts():
        filename = part.get_filename()
        if not filename:
            continue
        files.append((part.get_param("name", header="content-disposition") or "files", filename, part.get_payload(decode=True)))
    return files


def safe_extract_zip(zip_path: Path, destination: Path) -> None:
    with zipfile.ZipFile(zip_path) as archive:
        for member in archive.infolist():
            if member.is_dir():
                continue
            member_path = Path(member.filename)
            if member_path.is_absolute() or ".." in member_path.parts:
                continue
            target = destination / safe_name(member_path.name)
            target = unique_path(target)
            with archive.open(member) as source, target.open("wb") as output:
                shutil.copyfileobj(source, output)


def image_dimensions(path: Path) -> tuple[int, int] | None:
    try:
        with Image.open(path) as image:
            return image.size
    except Exception:
        return None


def analyze_session(session_id: str) -> dict:
    session = SESSION_DIR / session_id
    incoming = session / "incoming"
    expanded = session / "images"
    expanded.mkdir(parents=True, exist_ok=True)
    for path in incoming.iterdir():
        if path.suffix.lower() == ".zip":
            safe_extract_zip(path, expanded)
        elif path.suffix.lower() in IMAGE_EXTENSIONS:
            shutil.copy2(path, unique_path(expanded / safe_name(path.name)))

    records = []
    for index, path in enumerate(sorted(expanded.iterdir(), key=lambda p: p.name.lower()), start=1):
        if path.suffix.lower() not in IMAGE_EXTENSIONS:
            continue
        ocr_text = run_ocr(path)
        parsed = parse_address(f"{path.stem} {ocr_text}")
        number = photo_number(path.name, ocr_text)
        records.append({
            "id": hashlib.sha1(path.name.encode("utf-8")).hexdigest()[:12],
            "filename": path.name,
            "photo_number": number if number is not None else index,
            "ocr_text": ocr_text,
            "dimensions": image_dimensions(path),
            **parsed,
        })
    return {"session_id": session_id, "records": records, "ocr_available": True}


def set_run_font(run, *, blue: bool = False, size: int = 18) -> None:
    run.font.name = "BiauKai"
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "標楷體")
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "BiauKai")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "BiauKai")
    if blue:
        run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)


def set_cell_text(cell, text: str, *, blue: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(str(text))
    set_run_font(run, blue=blue)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def mark_row_keep_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def remove_table_rows_from(table, start: int) -> None:
    for row in list(table.rows)[start:]:
        table._tbl.remove(row._tr)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_image_to_cell(cell, image_path: Path) -> None:
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        with Image.open(image_path) as image:
            width, height = image.size
        target_height_cm = 9.0
        target_width_cm = target_height_cm * width / height
        if target_width_cm > 16.2:
            run.add_picture(str(image_path), width=Cm(16.2))
        else:
            run.add_picture(str(image_path), height=Cm(target_height_cm))
    except Exception:
        run.add_picture(str(image_path), height=Cm(9))


def rebuild_template_pages(doc: Document, page_count: int) -> None:
    body = doc.element.body
    children = list(body)
    paragraphs = [child for child in children if child.tag == qn("w:p") and "鑑定報告書現況照片" in "".join(child.itertext())]
    tables = [child for child in children if child.tag == qn("w:tbl")]
    if not paragraphs or not tables:
        raise RuntimeError("Word 範本結構不符：找不到標題或照片表格")
    first_title = deepcopy(paragraphs[0])
    later_title = deepcopy(paragraphs[1] if len(paragraphs) > 1 else paragraphs[0])
    base_table = deepcopy(tables[0])
    section = body.sectPr
    for child in list(body):
        if child is not section:
            body.remove(child)
    for page in range(page_count):
        body.insert(len(body) - 1, deepcopy(first_title if page == 0 else later_title))
        body.insert(len(body) - 1, deepcopy(base_table))


def build_word(address_name: str, images: list[tuple[dict, Path]], output_path: Path) -> None:
    doc = Document(BLANK_TEMPLATE)
    page_count = max(1, (len(images) + 1) // 2)
    rebuild_template_pages(doc, page_count)

    for paragraph in doc.paragraphs:
        if "鑑定報告書現況照片" in paragraph.text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=18)

    current = 0
    for table_index, table in enumerate(doc.tables):
        for row in table.rows:
            mark_row_keep_together(row)
        blocks = [(0, 1, 2, 3), (5, 6, 7, 8)]
        for header_row, number_row, number_row_2, image_row in blocks:
            if current >= len(images):
                remove_table_rows_from(table, 4 if header_row == 5 else 0)
                break
            record, image_path = images[current]
            number = current + 1
            set_cell_text(table.rows[header_row].cells[2], "A", blue=True)
            set_cell_text(table.rows[number_row].cells[0], number, blue=True)
            set_cell_text(table.rows[number_row_2].cells[0], number, blue=True)
            add_image_to_cell(table.rows[image_row].cells[0], image_path)
            current += 1
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def generate_results(session_id: str, records: list[dict]) -> Path:
    session = SESSION_DIR / session_id
    source_dir = session / "images"
    result_dir = session / "result"
    if result_dir.exists():
        shutil.rmtree(result_dir)
    result_dir.mkdir(parents=True)

    groups: dict[tuple[str, str], list[tuple[dict, Path]]] = {}
    pending_count = 0
    for record in records:
        source = source_dir / safe_name(record.get("filename", ""))
        if not source.exists():
            continue
        status = record.get("status", "incomplete")
        first = safe_name(record.get("first_level", ""), "地址不完整")
        second = safe_name(record.get("second_level", ""), "地址不完整")
        if status == "success" and record.get("first_level") and record.get("second_level"):
            groups.setdefault((first, second), []).append((record, source))
        else:
            pending_count += 1
            reason = STATUS_LABELS.get(status, "地址不完整")
            destination = result_dir / "待確認" / reason
            destination.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source, unique_path(destination / source.name))

    report_rows = []
    for (first, second), items in sorted(groups.items()):
        items.sort(key=lambda item: (int(item[0].get("photo_number") or 999999), item[0].get("filename", "")))
        address_dir = result_dir / first / second
        address_dir.mkdir(parents=True, exist_ok=True)
        copied_items = []
        for index, (record, source) in enumerate(items, start=1):
            new_name = f"照片{index:02d}_{safe_name(source.name)}"
            target = unique_path(address_dir / new_name)
            shutil.copy2(source, target)
            copied_items.append((record, target))
        word_name = f"{safe_name(first + second)}_貼照片.docx"
        build_word(f"{first}{second}", copied_items, address_dir / word_name)
        report_rows.append((first, second, len(items)))

    total = sum(count for _, _, count in report_rows) + pending_count
    with (result_dir / "分類報告.csv").open("w", encoding="utf-8-sig", newline="") as output:
        writer = csv.writer(output)
        writer.writerow(["分類報告", datetime.now().strftime("%Y-%m-%d %H:%M")])
        writer.writerow(["總照片數", total])
        writer.writerow(["成功分類數", total - pending_count])
        writer.writerow(["待確認數", pending_count])
        writer.writerow([])
        writer.writerow(["第一層地址", "第二層地址", "照片數量"])
        writer.writerows(report_rows)

    archive_path = session / "地址照片與Word成果.zip"
    if archive_path.exists():
        archive_path.unlink()
    with zipfile.ZipFile(archive_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in result_dir.rglob("*"):
            if path.is_file():
                archive.write(path, path.relative_to(result_dir))
    return archive_path


class Handler(BaseHTTPRequestHandler):
    server_version = "AddressPhotoTool/1.0"

    def log_message(self, format: str, *args) -> None:
        sys.stdout.write(f"[{self.log_date_time_string()}] {format % args}\n")

    def send_json(self, payload: dict, status: int = 200) -> None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def send_file(self, path: Path, download_name: str | None = None) -> None:
        if not path.exists() or not path.is_file():
            self.send_error(404)
            return
        content_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(path.stat().st_size))
        if download_name:
            self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{quote(download_name)}")
        self.end_headers()
        with path.open("rb") as source:
            shutil.copyfileobj(source, self.wfile)

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/health":
            self.send_json({"status": "ok"})
            return
        if parsed.path == "/":
            self.send_file(STATIC_DIR / "index.html")
            return
        if parsed.path.startswith("/api/preview/"):
            parts = parsed.path.split("/", 4)
            if len(parts) != 5:
                self.send_error(404)
                return
            session_id = safe_name(parts[3])
            filename = safe_name(unquote(parts[4]))
            self.send_file(SESSION_DIR / session_id / "images" / filename)
            return
        if parsed.path.startswith("/api/download/"):
            session_id = safe_name(parsed.path.rsplit("/", 1)[-1])
            self.send_file(SESSION_DIR / session_id / "地址照片與Word成果.zip", "地址照片與Word成果.zip")
            return
        self.send_error(404)

    def do_POST(self) -> None:
        try:
            if self.path == "/api/analyze":
                cleanup_expired_sessions()
                length = int(self.headers.get("Content-Length", "0"))
                if length <= 0 or length > MAX_UPLOAD_BYTES:
                    self.send_json({"error": "檔案大小不正確或超過服務限制"}, 400)
                    return
                files = parse_multipart(self.headers, self.rfile.read(length))
                if not files:
                    self.send_json({"error": "沒有收到照片或 ZIP"}, 400)
                    return
                session_id = uuid.uuid4().hex
                incoming = SESSION_DIR / session_id / "incoming"
                incoming.mkdir(parents=True)
                for _, filename, content in files:
                    target = unique_path(incoming / safe_name(Path(filename).name))
                    target.write_bytes(content)
                result = analyze_session(session_id)
                self.send_json(result)
                return
            if self.path == "/api/generate":
                length = int(self.headers.get("Content-Length", "0"))
                payload = json.loads(self.rfile.read(length).decode("utf-8"))
                session_id = safe_name(payload.get("session_id", ""))
                if not (SESSION_DIR / session_id).exists():
                    self.send_json({"error": "工作階段已失效，請重新上傳"}, 404)
                    return
                archive = generate_results(session_id, payload.get("records", []))
                self.send_json({"download_url": f"/api/download/{session_id}", "size": archive.stat().st_size})
                return
            self.send_error(404)
        except Exception as exc:
            self.send_json({"error": f"處理失敗：{exc}"}, 500)


def main() -> None:
    ensure_dirs()
    cleanup_expired_sessions()
    if not BLANK_TEMPLATE.exists():
        raise SystemExit(f"缺少 Word 範本：{BLANK_TEMPLATE}")
    host = os.environ.get("HOST", "0.0.0.0")
    port = int(os.environ.get("PORT", "8765"))
    server = ThreadingHTTPServer((host, port), Handler)
    print(f"照片分類與 Word 雲端服務已啟動：{host}:{port}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\n工具已停止")


if __name__ == "__main__":
    main()
